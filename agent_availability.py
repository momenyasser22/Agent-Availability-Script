#!/usr/bin/env python3
"""
Agent Availability Script - Interactive Terminal Application
Calculates agent availability per Domain and Operating System using CSV inputs and SQLite database.
"""

import csv
import sqlite3
import sys
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Tuple, Set, Optional
import os

import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl.workbook import Workbook

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Constants
WINDOWS_TABLE = "windows_agents"
LINUX_TABLE = "linux_agents"

# Column names for baseline CSVs
BASELINE_COLUMNS = ["Domain", "Agent Name"]

# Column names for availability CSVs
AVAILABILITY_COLUMNS = ["Domain", "Agent Name", "Last Available Date"]

# Red fill for unavailable hosts in Excel
RED_FILL = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
HEADER_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
HEADER_FONT = Font(bold=True, color="FFFFFF")

# Persistent data directory
DATA_DIR = Path("data")
REPORTS_DIR = Path("reports")
DB_FILE = DATA_DIR / "agent_baseline.db"


def get_data_dir() -> Path:
    """
    Get or create the data directory for persistent storage.

    Returns:
        Path to the data directory.
    """
    DATA_DIR.mkdir(exist_ok=True)
    return DATA_DIR


def get_reports_dir() -> Path:
    """
    Get or create the reports directory for output files.

    Returns:
        Path to the reports directory.
    """
    REPORTS_DIR.mkdir(exist_ok=True)
    return REPORTS_DIR


def get_db_path() -> Path:
    """
    Get the path to the persistent database file.

    Returns:
        Path to the database file.
    """
    get_data_dir()
    return DB_FILE


def show_database_info():
    """Display information about the persistent database."""
    db_path = get_db_path()

    if db_path.exists():
        # Get database size
        size_bytes = db_path.stat().st_size
        size_kb = size_bytes / 1024

        # Get file modification time
        mtime = datetime.fromtimestamp(db_path.stat().st_mtime)

        # Get agent counts
        conn = sqlite3.connect(str(db_path))
        cursor = conn.cursor()

        cursor.execute(f'SELECT COUNT(*) FROM {WINDOWS_TABLE}')
        windows_count = cursor.fetchone()[0]

        cursor.execute(f'SELECT COUNT(*) FROM {LINUX_TABLE}')
        linux_count = cursor.fetchone()[0]

        conn.close()

        print("\n" + "=" * 50)
        print("PERSISTENT DATABASE INFO")
        print("=" * 50)
        print(f"Location: {db_path.absolute()}")
        print(f"Size: {size_kb:.2f} KB")
        print(f"Last Modified: {mtime.strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"Windows Agents: {windows_count}")
        print(f"Linux Agents: {linux_count}")
        print("=" * 50)
    else:
        print("\n" + "=" * 50)
        print("PERSISTENT DATABASE INFO")
        print("=" * 50)
        print(f"Location: {db_path.absolute()}")
        print("Status: No baseline data loaded yet")
        print("=" * 50)


class AgentAvailabilityError(Exception):
    """Base exception for agent availability script errors."""
    pass


class CSVValidationError(AgentAvailabilityError):
    """Exception raised when CSV validation fails."""
    pass


def validate_baseline_csv(csv_path: Path) -> List[Tuple[str, str]]:
    """
    Validate and parse baseline CSV file.

    Args:
        csv_path: Path to the baseline CSV file.

    Returns:
        List of tuples containing (domain, agent_name).

    Raises:
        CSVValidationError: If CSV format is invalid.
    """
    if not csv_path.exists():
        raise CSVValidationError(f"CSV file not found: {csv_path}")

    if csv_path.stat().st_size == 0:
        raise CSVValidationError(f"CSV file is empty: {csv_path}")

    records = []

    with open(csv_path, 'r', encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)

        # Validate headers
        headers = reader.fieldnames
        if headers is None:
            raise CSVValidationError(f"CSV file has no headers: {csv_path}")

        # Strip BOM and whitespace from headers
        normalized_headers = [h.strip().strip('\ufeff') for h in headers]
        if normalized_headers != BASELINE_COLUMNS:
            raise CSVValidationError(
                f"Invalid headers in {csv_path}. Expected: {BASELINE_COLUMNS}, Got: {normalized_headers}"
            )

        # Parse rows
        for row_num, row in enumerate(reader, start=2):
            domain = row["Domain"].strip()
            agent_name = row["Agent Name"].strip()

            if not domain or not agent_name:
                raise CSVValidationError(
                    f"Empty domain or agent name in {csv_path} at row {row_num}"
                )

            records.append((domain, agent_name))

    if not records:
        raise CSVValidationError(f"No data records found in CSV: {csv_path}")

    return records


def validate_availability_csv(csv_path: Path) -> List[Tuple[str, str, str]]:
    """
    Validate and parse availability CSV file.

    Handles date fields that contain commas (e.g., "Jan 31, 2026 @ 12:38:00.504")

    Args:
        csv_path: Path to the availability CSV file.

    Returns:
        List of tuples containing (domain, agent_name, available_date_str).

    Raises:
        CSVValidationError: If CSV format is invalid.
    """
    if not csv_path.exists():
        raise CSVValidationError(f"CSV file not found: {csv_path}")

    if csv_path.stat().st_size == 0:
        # Empty availability CSV is valid - just no available agents
        return []

    records = []

    with open(csv_path, 'r', encoding='utf-8-sig') as f:
        # Read all lines
        lines = f.readlines()

        if not lines:
            raise CSVValidationError(f"CSV file is empty: {csv_path}")

        # Parse header - handle quoted column names
        header_line = lines[0].strip()
        headers = [h.strip().strip('\ufeff').strip('"').strip("'") for h in header_line.split(',')]

        # Validate headers - case-insensitive comparison
        headers_normalized = [h.strip().strip('\ufeff').strip('"').strip("'").lower() for h in headers]
        expected_normalized = [h.lower() for h in AVAILABILITY_COLUMNS]
        if headers_normalized != expected_normalized:
            raise CSVValidationError(
                f"Invalid headers in {csv_path}. Expected: {AVAILABILITY_COLUMNS}, Got: {headers}"
            )

        # Parse data rows
        for row_num, line in enumerate(lines[1:], start=2):
            line = line.strip()
            if not line:
                continue

            # Handle quoted/unquoted CSV values
            # Split by comma but respect quotes
            parts = []
            current = ""
            in_quotes = False

            for char in line:
                if char == '"':
                    in_quotes = not in_quotes
                elif char == ',' and not in_quotes:
                    parts.append(current)
                    current = ""
                else:
                    current += char
            parts.append(current)

            # Ensure we have at least 3 parts
            if len(parts) < 3:
                raise CSVValidationError(
                    f"Invalid row format in {csv_path} at row {row_num}"
                )

            # First two columns are Domain and Agent Name
            # Strip quotes and whitespace from values
            domain = parts[0].strip().strip('"').strip("'")
            agent_name = parts[1].strip().strip('"').strip("'")

            # The date column is everything from index 2 onwards, joined with commas
            # (because dates like "Jan 31, 2026 @ 12:38:00.504" get split on the comma)
            available_date = ", ".join(p.strip().strip('"').strip("'") for p in parts[2:])

            if not agent_name or not domain:
                raise CSVValidationError(
                    f"Empty domain or agent name in {csv_path} at row {row_num}"
                )

            records.append((domain, agent_name, available_date))

    if not records:
        return []

    return records


def create_database():
    """
    Create SQLite database with required tables.

    Creates:
        - windows_agents table
        - linux_agents table
    """
    conn = sqlite3.connect(str(get_db_path()))
    cursor = conn.cursor()

    # Create Windows agents table
    cursor.execute(f'''
        CREATE TABLE IF NOT EXISTS {WINDOWS_TABLE} (
            domain TEXT NOT NULL,
            agent_name TEXT NOT NULL,
            operating_system TEXT NOT NULL DEFAULT 'Windows',
            UNIQUE(domain, agent_name)
        )
    ''')

    # Create Linux agents table
    cursor.execute(f'''
        CREATE TABLE IF NOT EXISTS {LINUX_TABLE} (
            domain TEXT NOT NULL,
            agent_name TEXT NOT NULL,
            operating_system TEXT NOT NULL DEFAULT 'Linux',
            UNIQUE(domain, agent_name)
        )
    ''')

    conn.commit()
    conn.close()


def load_baseline_windows(csv_path: Path):
    """
    Load Windows baseline data from CSV into database.

    Args:
        csv_path: Path to the Windows baseline CSV file.

    Action:
        - Clears existing Windows data
        - Inserts all records from CSV
    """
    records = validate_baseline_csv(csv_path)

    conn = sqlite3.connect(str(get_db_path()))
    cursor = conn.cursor()

    # Clear existing Windows data
    cursor.execute(f'DELETE FROM {WINDOWS_TABLE}')

    # Insert new records
    for domain, agent_name in records:
        cursor.execute(
            f'INSERT INTO {WINDOWS_TABLE} (domain, agent_name, operating_system) VALUES (?, ?, ?)',
            (domain, agent_name, 'Windows')
        )

    conn.commit()
    conn.close()

    print(f"\nSuccessfully loaded {len(records)} Windows agents from {csv_path}")


def load_baseline_linux(csv_path: Path):
    """
    Load Linux baseline data from CSV into database.

    Args:
        csv_path: Path to the Linux baseline CSV file.

    Action:
        - Clears existing Linux data
        - Inserts all records from CSV
    """
    records = validate_baseline_csv(csv_path)

    conn = sqlite3.connect(str(get_db_path()))
    cursor = conn.cursor()

    # Clear existing Linux data
    cursor.execute(f'DELETE FROM {LINUX_TABLE}')

    # Insert new records
    for domain, agent_name in records:
        cursor.execute(
            f'INSERT INTO {LINUX_TABLE} (domain, agent_name, operating_system) VALUES (?, ?, ?)',
            (domain, agent_name, 'Linux')
        )

    conn.commit()
    conn.close()

    print(f"\nSuccessfully loaded {len(records)} Linux agents from {csv_path}")


def parse_available_date(date_str: str) -> datetime:
    """
    Parse date string from CSV.

    Supports multiple date formats:
    - 'YYYY-MM-DD HH:MM:SS'
    - 'MMM DD, YYYY @ HH:MM:SS.fff' (e.g., 'Jan 31, 2026 @ 12:38:00.504')

    Args:
        date_str: Date string in various formats.

    Returns:
        datetime object.

    Raises:
        CSVValidationError: If date format is invalid.
    """
    # List of supported date formats
    date_formats = [
        "%Y-%m-%d %H:%M:%S",           # Standard format: 2026-01-31 12:38:00
        "%b %d, %Y @ %H:%M:%S.%f",     # User's format: Jan 31, 2026 @ 12:38:00.504
        "%b %d, %Y @ %H:%M:%S",        # Without milliseconds: Jan 31, 2026 @ 12:38:00
        "%B %d, %Y @ %H:%M:%S.%f",     # Full month name: January 31, 2026 @ 12:38:00.504
        "%B %d, %Y @ %H:%M:%S",        # Full month name without milliseconds
        "%d-%m-%Y %H:%M:%S",           # DD-MM-YYYY HH:MM:SS
        "%m/%d/%Y %H:%M:%S",           # MM/DD/YYYY HH:MM:SS
        "%Y/%m/%d %H:%M:%S",           # YYYY/MM/DD HH:MM:SS
    ]

    # Clean the date string (remove extra whitespace)
    date_str = date_str.strip()

    for fmt in date_formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue

    # If none of the formats worked, raise error
    raise CSVValidationError(
        f"Invalid date format: '{date_str}'. "
        f"Supported formats include: 'YYYY-MM-DD HH:MM:SS', 'Jan 31, 2026 @ 12:38:00.504'"
    )


def is_within_last_24_hours(available_date: datetime) -> bool:
    """
    Check if a datetime is within the last 24 hours from now.

    Args:
        available_date: The datetime to check.

    Returns:
        True if within last 24 hours, False otherwise.
    """
    now = datetime.now()
    twenty_four_hours_ago = now - timedelta(hours=24)
    return available_date >= twenty_four_hours_ago


def get_baseline_agents(table_name: str) -> Dict[str, Set[str]]:
    """
    Get all baseline agents from database grouped by domain.

    Args:
        table_name: Name of the table to query.

    Returns:
        Dictionary mapping domain to set of agent names.
    """
    conn = sqlite3.connect(str(get_db_path()))
    cursor = conn.cursor()

    cursor.execute(f'SELECT DISTINCT domain, agent_name FROM {table_name}')
    rows = cursor.fetchall()

    conn.close()

    agents_by_domain: Dict[str, Set[str]] = {}

    for domain, agent_name in rows:
        if domain not in agents_by_domain:
            agents_by_domain[domain] = set()
        agents_by_domain[domain].add(agent_name)

    return agents_by_domain


def get_availability_records(csv_path: Path) -> Dict[Tuple[str, str], datetime]:
    """
    Parse availability CSV and return records grouped by (agent_name, domain).

    Args:
        csv_path: Path to the availability CSV file.

    Returns:
        Dictionary mapping (agent_name, domain) to available_date.
    """
    records = validate_availability_csv(csv_path)

    availability_map: Dict[Tuple[str, str], datetime] = {}

    # records are now (domain, agent_name, available_date_str)
    for domain, agent_name, available_date_str in records:
        available_date = parse_available_date(available_date_str)
        key = (agent_name, domain)
        # Keep the most recent date if duplicates exist
        if key not in availability_map or available_date > availability_map[key]:
            availability_map[key] = available_date

    return availability_map


def calculate_availability(
    baseline_agents: Dict[str, Set[str]],
    availability_map: Dict[Tuple[str, str], datetime]
) -> Dict[str, Dict[str, Dict]]:
    """
    Calculate agent availability per domain.

    An agent is AVAILABLE only if:
        1. It exists in the baseline database
        2. It exists in the availability CSV
        3. Available Date is within the last 24 hours

    Args:
        baseline_agents: Dictionary mapping domain to set of agent names from baseline.
        availability_map: Dictionary mapping (agent_name, domain) to available_date.

    Returns:
        Dictionary with structure:
        {
            "domain_name": {
                "total": int,
                "available": set of agent names,
                "not_available": set of agent names,
                "last_available_dates": dict mapping agent_name to last available date
            }
        }
    """
    results = {}

    for domain, agents in baseline_agents.items():
        domain_results = {
            "total": len(agents),
            "available": set(),
            "not_available": set(),
            "last_available_dates": {}
        }

        for agent_name in agents:
            key = (agent_name, domain)

            # Check rule 1: exists in baseline (always true here)
            # Check rule 2: exists in availability CSV
            if key not in availability_map:
                domain_results["not_available"].add(agent_name)
                continue

            # Check rule 3: within last 24 hours
            available_date = availability_map[key]
            if is_within_last_24_hours(available_date):
                domain_results["available"].add(agent_name)
                domain_results["last_available_dates"][agent_name] = available_date
            else:
                domain_results["not_available"].add(agent_name)
                domain_results["last_available_dates"][agent_name] = available_date

        results[domain] = domain_results

    return results


def print_console_report(windows_results: Dict, linux_results: Dict):
    """
    Print availability report to console.

    Format:
        Operating System: <OS>
        Domain: <Domain>
        Hosts Not Available:
            * HOST01
            * HOST02
        Availability Percentage: XX%
    """
    print("\n" + "=" * 50)
    print("AGENT AVAILABILITY REPORT")
    print("=" * 50 + "\n")

    # Windows domains
    for domain in sorted(windows_results.keys()):
        result = windows_results[domain]
        total = result["total"]
        available_count = len(result["available"])
        not_available = result["not_available"]

        print(f"Operating System: Windows")
        print(f"Domain: {domain}")

        if not_available:
            print("Hosts Not Available:")
            for host in sorted(not_available):
                print(f"\t*\t{host}")
        else:
            print("Hosts Not Available: None")

        percentage = (available_count / total * 100) if total > 0 else 0
        print(f"Availability Percentage: {percentage:.1f}%")
        print()

    # Linux domains
    for domain in sorted(linux_results.keys()):
        result = linux_results[domain]
        total = result["total"]
        available_count = len(result["available"])
        not_available = result["not_available"]

        print(f"Operating System: Linux")
        print(f"Domain: {domain}")

        if not_available:
            print("Hosts Not Available:")
            for host in sorted(not_available):
                print(f"\t*\t{host}")
        else:
            print("Hosts Not Available: None")

        percentage = (available_count / total * 100) if total > 0 else 0
        print(f"Availability Percentage: {percentage:.1f}%")
        print()

    print("=" * 50)


def set_cell_border(cell):
    """Set borders for a cell in docx table."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)

    tcPr.append(tcBorders)


def generate_xlsx_report(
    windows_results: Dict,
    linux_results: Dict,
    output_path: Path
):
    """
    Generate Excel (.xlsx) report with Windows and Linux sheets.

    Columns:
        - OS
        - Domain
        - Agent Name
        - Status
        - Last Available Date

    Rows with NOT AVAILABLE hosts are highlighted RED.
    """
    wb = Workbook()

    # Remove default sheet
    if "Sheet" in wb.sheetnames:
        del wb["Sheet"]

    # Define column headers
    headers = ["OS", "Domain", "Agent Name", "Status", "Last Available Date"]

    # Create Windows sheet
    ws_windows = wb.create_sheet("Windows")
    ws_windows.append(headers)

    # Format header row
    for cell in ws_windows[1]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal='center', vertical='center')

    for domain, result in windows_results.items():
        # Available agents
        for agent in sorted(result["available"]):
            last_date = result["last_available_dates"][agent]
            row = [
                "Windows",
                domain,
                agent,
                "Available",
                last_date.strftime("%Y-%m-%d %H:%M:%S")
            ]
            ws_windows.append(row)

        # Not available agents
        for agent in sorted(result["not_available"]):
            last_date = result["last_available_dates"].get(agent)
            date_str = last_date.strftime("%Y-%m-%d %H:%M:%S") if last_date else "N/A"
            row = [
                "Windows",
                domain,
                agent,
                "Not Available",
                date_str
            ]
            row_idx = ws_windows.max_row
            ws_windows.append(row)
            # Highlight row in red
            for cell in ws_windows[row_idx]:
                cell.fill = RED_FILL

    # Auto-adjust column widths
    for column in ws_windows.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws_windows.column_dimensions[column_letter].width = adjusted_width

    # Create Linux sheet
    ws_linux = wb.create_sheet("Linux")
    ws_linux.append(headers)

    # Format header row
    for cell in ws_linux[1]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal='center', vertical='center')

    for domain, result in linux_results.items():
        # Available agents
        for agent in sorted(result["available"]):
            last_date = result["last_available_dates"][agent]
            row = [
                "Linux",
                domain,
                agent,
                "Available",
                last_date.strftime("%Y-%m-%d %H:%M:%S")
            ]
            ws_linux.append(row)

        # Not available agents
        for agent in sorted(result["not_available"]):
            last_date = result["last_available_dates"].get(agent)
            date_str = last_date.strftime("%Y-%m-%d %H:%M:%S") if last_date else "N/A"
            row = [
                "Linux",
                domain,
                agent,
                "Not Available",
                date_str
            ]
            row_idx = ws_linux.max_row
            ws_linux.append(row)
            # Highlight row in red
            for cell in ws_linux[row_idx]:
                cell.fill = RED_FILL

    # Auto-adjust column widths
    for column in ws_linux.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws_linux.column_dimensions[column_letter].width = adjusted_width

    wb.save(output_path)
    print(f"\nExcel report generated: {output_path}")


def generate_docx_report(
    windows_results: Dict,
    linux_results: Dict,
    output_path: Path
):
    """
    Generate Word (.docx) report with availability information.

    Format for EACH OS and Domain:
        Heading: OS + Domain
        Table of unavailable hosts
        Availability percentage clearly stated

    Example structure:
        Windows - Domain1
        Unavailable Hosts:
        | Agent Name | Last Available Date |
        Availability: 75%
    """
    doc = Document()

    # Title
    title = doc.add_heading("AGENT AVAILABILITY REPORT", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Report generation time
    timestamp = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # Blank line

    # Process Windows domains
    if windows_results:
        for domain in sorted(windows_results.keys()):
            result = windows_results[domain]
            total = result["total"]
            available_count = len(result["available"])
            not_available = result["not_available"]
            percentage = (available_count / total * 100) if total > 0 else 0

            # Heading: OS + Domain
            heading_text = f"Windows - {domain}"
            doc.add_heading(heading_text, level=2)

            # Table of unavailable hosts
            doc.add_paragraph("Unavailable Hosts:", style='Heading 3')

            if not_available:
                # Create table with headers
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'

                # Set headers
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Agent Name"
                hdr_cells[1].text = "Last Available Date"

                # Format header cells
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    set_cell_border(cell)

                # Add unavailable hosts to table
                for agent in sorted(not_available):
                    last_date = result["last_available_dates"].get(agent)
                    date_str = last_date.strftime("%Y-%m-%d %H:%M:%S") if last_date else "N/A"

                    row_cells = table.add_row().cells
                    row_cells[0].text = agent
                    row_cells[1].text = date_str

                    for cell in row_cells:
                        set_cell_border(cell)
            else:
                doc.add_paragraph("No unavailable hosts.")

            # Availability percentage
            p = doc.add_paragraph()
            p.add_run("Availability: ").bold = True

            percentage_text = f"{percentage:.1f}%"
            run = p.add_run(percentage_text)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00) if percentage >= 75 else RGBColor(0xFF, 0x00, 0x00)

            doc.add_paragraph()  # Blank line between domains

    # Process Linux domains
    if linux_results:
        for domain in sorted(linux_results.keys()):
            result = linux_results[domain]
            total = result["total"]
            available_count = len(result["available"])
            not_available = result["not_available"]
            percentage = (available_count / total * 100) if total > 0 else 0

            # Heading: OS + Domain
            heading_text = f"Linux - {domain}"
            doc.add_heading(heading_text, level=2)

            # Table of unavailable hosts
            doc.add_paragraph("Unavailable Hosts:", style='Heading 3')

            if not_available:
                # Create table with headers
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'

                # Set headers
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Agent Name"
                hdr_cells[1].text = "Last Available Date"

                # Format header cells
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    set_cell_border(cell)

                # Add unavailable hosts to table
                for agent in sorted(not_available):
                    last_date = result["last_available_dates"].get(agent)
                    date_str = last_date.strftime("%Y-%m-%d %H:%M:%S") if last_date else "N/A"

                    row_cells = table.add_row().cells
                    row_cells[0].text = agent
                    row_cells[1].text = date_str

                    for cell in row_cells:
                        set_cell_border(cell)
            else:
                doc.add_paragraph("No unavailable hosts.")

            # Availability percentage
            p = doc.add_paragraph()
            p.add_run("Availability: ").bold = True

            percentage_text = f"{percentage:.1f}%"
            run = p.add_run(percentage_text)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x00, 0xFF, 0x00) if percentage >= 75 else RGBColor(0xFF, 0x00, 0x00)

            doc.add_paragraph()  # Blank line between domains

    doc.save(output_path)
    print(f"Word report generated: {output_path}")


def check_availability(windows_csv: Path, linux_csv: Path, output_base: str):
    """
    Main function to check agent availability.

    Args:
        windows_csv: Path to Windows availability CSV.
        linux_csv: Path to Linux availability CSV.
        output_base: Base name for output reports (without extension).
    """
    # Get baseline agents from database
    windows_baseline = get_baseline_agents(WINDOWS_TABLE)
    linux_baseline = get_baseline_agents(LINUX_TABLE)

    if not windows_baseline and not linux_baseline:
        print("\nError: No baseline data found in database.")
        print("Please load baseline data first using options 1 or 2.")
        return False

    # Parse availability CSVs
    windows_availability = get_availability_records(windows_csv)
    linux_availability = get_availability_records(linux_csv)

    # Calculate availability
    windows_results = calculate_availability(windows_baseline, windows_availability)
    linux_results = calculate_availability(linux_baseline, linux_availability)

    # Print console report
    print_console_report(windows_results, linux_results)

    # Generate reports in the reports directory
    reports_dir = get_reports_dir()
    xlsx_path = reports_dir / f"{output_base}.xlsx"
    docx_path = reports_dir / f"{output_base}.docx"

    generate_xlsx_report(windows_results, linux_results, xlsx_path)
    generate_docx_report(windows_results, linux_results, docx_path)

    print(f"\nReports saved to: {reports_dir.absolute()}")

    return True


def get_input_path(prompt: str) -> Optional[Path]:
    """
    Get file path from user input with validation.

    Args:
        prompt: The prompt to display to the user.

    Returns:
        Path object if valid file exists, None if user wants to cancel.
    """
    while True:
        user_input = input(prompt).strip()

        if not user_input:
            return None

        if user_input.lower() in ('quit', 'exit', 'cancel'):
            return None

        path = Path(user_input)

        if path.exists():
            return path
        else:
            print(f"Error: File not found: {path}")
            retry = input("Try again? (y/n): ").strip().lower()
            if retry != 'y':
                return None


def display_main_menu():
    """Display the main menu options."""
    print("\n" + "=" * 60)
    print(" "*15 + "AGENT AVAILABILITY SYSTEM")
    print("=" * 60)
    print("\nMain Menu:")
    print("\t[1] Load Windows Baseline")
    print("\t[2] Load Linux Baseline")
    print("\t[3] Check Availability & Generate Report")
    print("\t[4] View Database Info")
    print("\t[5] Exit")
    print("-" * 60)


def main():
    """Main interactive terminal application."""
    # Create database if needed
    create_database()

    # Show initial database info
    show_database_info()

    print("\nWelcome to Agent Availability System!")
    print("This tool calculates agent availability per Domain and Operating System.")

    while True:
        display_main_menu()
        choice = input("\nSelect an option [1-5]: ").strip()

        if choice == '1':
            # Load Windows Baseline
            print("\n--- Load Windows Baseline ---")
            csv_path = get_input_path("Enter Windows baseline CSV path (or 'cancel'): ")

            if csv_path:
                try:
                    load_baseline_windows(csv_path)
                except AgentAvailabilityError as e:
                    print(f"Error: {e}")
                except Exception as e:
                    print(f"Unexpected error: {e}")

            input("\nPress Enter to continue...")

        elif choice == '2':
            # Load Linux Baseline
            print("\n--- Load Linux Baseline ---")
            csv_path = get_input_path("Enter Linux baseline CSV path (or 'cancel'): ")

            if csv_path:
                try:
                    load_baseline_linux(csv_path)
                except AgentAvailabilityError as e:
                    print(f"Error: {e}")
                except Exception as e:
                    print(f"Unexpected error: {e}")

            input("\nPress Enter to continue...")

        elif choice == '3':
            # Check Availability
            print("\n--- Check Availability & Generate Report ---")

            windows_csv = get_input_path("Enter Windows availability CSV path (or 'cancel' if no Windows agents): ")
            linux_csv = get_input_path("Enter Linux availability CSV path (or 'cancel' if no Linux agents): ")

            if windows_csv is None and linux_csv is None:
                print("No CSV files provided. Operation cancelled.")
            else:
                output_name = input("\nEnter output report filename base (default: agent_availability_report): ").strip()
                if not output_name:
                    output_name = "agent_availability_report"

                # Remove extension if provided
                if output_name.endswith('.xlsx'):
                    output_name = output_name[:-5]
                elif output_name.endswith('.docx'):
                    output_name = output_name[:-5]

                try:
                    # Handle None values
                    if windows_csv is None:
                        windows_csv = Path("placeholder_empty.csv")
                        # Create empty file for validation
                        with open(windows_csv, 'w') as f:
                            f.write("Domain,Agent Name,Last Available Date\n")
                    if linux_csv is None:
                        linux_csv = Path("placeholder_empty.csv")
                        with open(linux_csv, 'w') as f:
                            f.write("Domain,Agent Name,Last Available Date\n")

                    check_availability(windows_csv, linux_csv, output_name)

                    # Clean up placeholders
                    if windows_csv.name.startswith("placeholder"):
                        windows_csv.unlink()
                    if linux_csv.name.startswith("placeholder"):
                        linux_csv.unlink()

                except AgentAvailabilityError as e:
                    print(f"Error: {e}")
                    if windows_csv.name.startswith("placeholder"):
                        windows_csv.unlink()
                    if linux_csv.name.startswith("placeholder"):
                        linux_csv.unlink()
                except Exception as e:
                    print(f"Unexpected error: {e}")
                    if windows_csv.name.startswith("placeholder"):
                        windows_csv.unlink()
                    if linux_csv.name.startswith("placeholder"):
                        linux_csv.unlink()

            input("\nPress Enter to continue...")

        elif choice == '4':
            # View Database Info
            show_database_info()
            input("\nPress Enter to continue...")

        elif choice == '5':
            # Exit
            print("\nThank you for using Agent Availability System!")
            print("Goodbye!")
            sys.exit(0)

        else:
            print("\nInvalid option. Please select 1-4.")
            input("Press Enter to continue...")


if __name__ == "__main__":
    main()
